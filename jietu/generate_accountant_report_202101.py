# -*- coding: utf-8 -*-
import csv
import json
from collections import Counter, OrderedDict
from pathlib import Path

import xlrd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from cli_anything.wukong.core.session import load_session
from cli_anything.wukong.utils.wukong_backend import WukongClient
from cli_anything.wukong.core import ledger, report, statement

ACCOUNT_ID = '2036268933673287680'
ACCOUNT_NAME = '幻式（江门新会）服装有限公司'
XLS_PATH = Path('/Users/czj/Desktop/kauiji/202101记账凭证.xls')
OUT_MD = Path('jietu/202101会计汇报-给会计看.md')
OUT_CSV = Path('jietu/202101做账结果清单.csv')
OUT_JSON = Path('jietu/202101会计汇报-给会计看.json')
OUT_DOCX = Path('output/doc/202101会计汇报-给会计看.docx')

SUBJECTS = {
    '1002 银行存款': 2036268933799116800,
    '1122 应收账款': 2036268933803311110,
    '1405 库存商品': 2036268933811699715,
    '2202 应付账款': 2036268933820088327,
    '3103 本年利润': 2036268933832671233,
    '5601 销售费用': 2036268933836865552,
    '5602 管理费用': 2036268933841059849,
}


def get_client():
    sess = load_session()
    client = WukongClient(sess['base_url'], token=sess['token'])
    client.post('/financeAccountSet/switchAccountSet', params={'accountId': ACCOUNT_ID})
    return client


def load_vouchers(path: Path):
    book = xlrd.open_workbook(str(path))
    sheet = book.sheet_by_name('凭证列表')
    groups = OrderedDict()
    for i in range(1, sheet.nrows):
        row = sheet.row_values(i)
        voucher_num = str(row[1]).strip()
        if not voucher_num:
            continue
        groups.setdefault(voucher_num, {
            'date': str(row[0]).strip(),
            'digest': str(row[2]).strip(),
            'rows': [],
        })
        groups[voucher_num]['rows'].append(row)
    result = []
    for voucher_num, item in groups.items():
        debit = round(sum(float(r[8] or 0) for r in item['rows']), 2)
        credit = round(sum(float(r[11] or 0) for r in item['rows']), 2)
        result.append({
            'date': item['date'],
            'voucher_num': voucher_num,
            'voucher_type': voucher_num.split('-')[0],
            'digest': item['digest'],
            'entry_count': len(item['rows']),
            'debit_total': debit,
            'credit_total': credit,
        })
    return result


def extract_detail_summary(client):
    out = {}
    for name, sid in SUBJECTS.items():
        rows = ledger.query_detail_account(client, sid, '2021-01-01', '2021-01-31')
        ending = rows[-1] if rows else {}
        out[name] = {
            'entry_count': len([r for r in rows if r.get('type') == '2']),
            'balance_direction': ending.get('balanceDirection'),
            'balance': ending.get('balance'),
            'debit_total': ending.get('debtorBalance'),
            'credit_total': ending.get('creditBalance'),
        }
    return out


def nonzero_rows(rows, amount_key):
    result = []
    for row in rows:
        value = row.get(amount_key)
        if isinstance(value, (int, float)) and abs(value) > 1e-9:
            result.append({'name': row.get('name'), amount_key: value})
    return result


def get_report_data(client):
    bs = report.balance_sheet(client, 1, '2021-01-31')
    income = report.income_statement(client, 1, '2021-01-31')
    cash = report.cash_flow_statement(client, 1, '2021-01-31')

    try:
        bs_check = report.balance_sheet_check(client, 1, '2021-01-31')
    except Exception as e:
        bs_check = {'error': str(e)}
    try:
        income_check = report.income_statement_check(client, 1, '2021-01-31')
    except Exception as e:
        income_check = {'error': str(e)}
    try:
        cash_check = report.cash_flow_check(client, 1, '2021-01-31')
    except Exception as e:
        cash_check = {'error': str(e)}

    return {
        'balance_sheet': {
            'rows': nonzero_rows(bs, 'endPeriod'),
            'check': bs_check,
        },
        'income_statement': {
            'rows': nonzero_rows(income, 'monthValue'),
            'check': income_check,
        },
        'cash_flow_statement': {
            'rows': nonzero_rows(cash, 'monthValue'),
            'check': cash_check,
        },
    }


def get_statement_status(client):
    data = statement.query_statement(client)
    return {
        'settleTime': data.get('settleTime'),
        'items': [
            {'name': item.get('statementName'), 'done': item.get('isEndOver')}
            for item in (data.get('statements') or [])
        ],
        'certificate_count_text': data.get('certificateNum'),
    }


def write_csv(vouchers):
    OUT_CSV.parent.mkdir(parents=True, exist_ok=True)
    with OUT_CSV.open('w', encoding='utf-8-sig', newline='') as f:
        writer = csv.DictWriter(f, fieldnames=['date', 'voucher_num', 'voucher_type', 'digest', 'entry_count', 'debit_total', 'credit_total'])
        writer.writeheader()
        writer.writerows(vouchers)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    run = p.add_run(text)
    run.bold = True
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    return table


def generate_docx(data):
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'PingFang SC'
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('2021年01月做账结果汇报')
    r.bold = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f'单位：{ACCOUNT_NAME}')

    doc.add_paragraph(f'本报告基于 {XLS_PATH.name} 已确认凭证结果，整理为会计复核用版本。')

    add_heading(doc, '一、做账结果概览', 1)
    overview_rows = [
        ['账套名称', ACCOUNT_NAME],
        ['账套ID', ACCOUNT_ID],
        ['做账期间', '2021年01月'],
        ['已整理凭证数', data['voucher_count']],
        ['凭证字分布', data['voucher_type_text']],
    ]
    add_table(doc, ['项目', '结果'], overview_rows)

    add_heading(doc, '二、关键科目结果', 1)
    k_rows = []
    for name, item in data['ledger_summary'].items():
        k_rows.append([name, item['entry_count'], item['balance_direction'], item['balance'], item['debit_total'], item['credit_total']])
    add_table(doc, ['科目', '业务笔数', '期末方向', '期末余额', '本期借方累计', '本期贷方累计'], k_rows)

    add_heading(doc, '三、资产负债表关键信息', 1)
    bs_rows = [[row['name'], row['endPeriod']] for row in data['reports']['balance_sheet']['rows']]
    add_table(doc, ['项目', '期末数'], bs_rows)

    add_heading(doc, '四、利润表关键信息', 1)
    income_rows = data['reports']['income_statement']['rows']
    if income_rows:
        add_table(doc, ['项目', '本月数'], [[row['name'], row['monthValue']] for row in income_rows])
    else:
        doc.add_paragraph('本次系统返回的利润表金额均为 0。')

    add_heading(doc, '五、现金流量表关键信息', 1)
    cash_rows = [[row['name'], row['monthValue']] for row in data['reports']['cash_flow_statement']['rows']]
    add_table(doc, ['项目', '本月数'], cash_rows)

    add_heading(doc, '六、需要会计重点复核的事项', 1)
    for item in data['review_points']:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '七、附件说明', 1)
    doc.add_paragraph(f'附件一：{OUT_CSV.name}（53张凭证结果清单）')

    doc.save(str(OUT_DOCX))


def main():
    client = get_client()
    vouchers = load_vouchers(XLS_PATH)
    voucher_counter = Counter(v['voucher_type'] for v in vouchers)
    ledger_summary = extract_detail_summary(client)
    reports = get_report_data(client)
    statement_status = get_statement_status(client)

    assets_total = next((x['endPeriod'] for x in reports['balance_sheet']['rows'] if x['name'] == '资产总计'), None)
    liability_equity_total = next((x['endPeriod'] for x in reports['balance_sheet']['rows'] if x['name'] == '负债和所有者权益（或股东权益）总计'), None)

    data = {
        'voucher_count': len(vouchers),
        'voucher_type_text': '，'.join(f'{k}{v}张' for k, v in voucher_counter.items()),
        'ledger_summary': ledger_summary,
        'reports': reports,
        'statement_status': statement_status,
        'review_points': [
            f'资产负债表显示资产总计 {assets_total}，负债和所有者权益总计 {liability_equity_total}；两者未平，需继续复核。',
            '利润表本次系统返回金额为 0，结合业务情况判断，当前数据更像“银行流水凭证结果”，可能并非完整收入成本结转链条。',
            f"应收账款期末余额显示为 {ledger_summary['1122 应收账款']['balance']}，呈异常贷方挂账特征，建议核对是否缺少销售确认或期初往来承接。",
            '结账模块（转出未交增值税、计提地税、计提所得税、结转损益）均未完成，不建议直接据此结账。',
        ],
    }

    write_csv(vouchers)
    OUT_JSON.write_text(json.dumps({'meta': data, 'vouchers': vouchers}, ensure_ascii=False, indent=2), encoding='utf-8')

    md = []
    md.append('# 2021年01月做账结果汇报')
    md.append('')
    md.append(f'- 单位：`{ACCOUNT_NAME}`')
    md.append(f'- 账套ID：`{ACCOUNT_ID}`')
    md.append('- 做账期间：`2021年01月`')
    md.append(f'- 数据基础：`{XLS_PATH.name}` 已确认凭证结果')
    md.append('')
    md.append('## 一、做账结果概览')
    md.append(f"- 本月已整理凭证共 `53` 张，其中 `{data['voucher_type_text']}`。")
    md.append('- 已形成凭证结果清单，便于会计逐张核对。')
    md.append('')
    md.append('## 二、关键科目结果')
    for name, item in ledger_summary.items():
        md.append(f"- `{name}`：期末 `{item['balance_direction']}` 方余额 `{item['balance']}`，本期借方累计 `{item['debit_total']}`，本期贷方累计 `{item['credit_total']}`。")
    md.append('')
    md.append('## 三、资产负债表关键信息')
    for row in reports['balance_sheet']['rows']:
        md.append(f"- `{row['name']}`：`{row['endPeriod']}`")
    md.append('')
    md.append('## 四、利润表关键信息')
    if reports['income_statement']['rows']:
        for row in reports['income_statement']['rows']:
            md.append(f"- `{row['name']}`：`{row['monthValue']}`")
    else:
        md.append('- 本次系统返回的利润表金额均为 `0`。')
    md.append('')
    md.append('## 五、现金流量表关键信息')
    for row in reports['cash_flow_statement']['rows']:
        md.append(f"- `{row['name']}`：`{row['monthValue']}`")
    md.append('')
    md.append('## 六、需要会计重点复核的事项')
    for item in data['review_points']:
        md.append(f'- {item}')
    md.append('')
    md.append('## 七、附件')
    md.append(f'- `jietu/202101做账结果清单.csv`：53张凭证结果清单')
    OUT_MD.write_text('\n'.join(md), encoding='utf-8')

    generate_docx(data)
    print(OUT_MD)
    print(OUT_CSV)
    print(OUT_JSON)
    print(OUT_DOCX)

if __name__ == '__main__':
    main()
