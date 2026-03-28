# -*- coding: utf-8 -*-
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from cli_anything.wukong.core.session import load_session
from cli_anything.wukong.utils.wukong_backend import WukongClient
from cli_anything.wukong.core import ledger, report

ACCOUNT_ID = '2036268933673287680'
ACCOUNT_NAME = '幻式（江门新会）服装有限公司'
PERIOD_TEXT = '2021年01月'
PERIOD_END = '2021-01-31'
START_DATE = '2021-01-01'
END_DATE = '2021-01-31'

OUT_DIR = Path('jietu')
DOCX_PATH = Path('output/doc/202101月度会计报表.docx')
JSON_PATH = OUT_DIR / '202101月度会计报表.json'
BALANCE_MD = OUT_DIR / '202101资产负债表.md'
INCOME_MD = OUT_DIR / '202101利润表.md'
CASH_MD = OUT_DIR / '202101现金流量表.md'
LEDGER_MD = OUT_DIR / '202101科目余额与关键科目.md'

KEY_SUBJECTS = {
    '1002 银行存款': 2036268933799116800,
    '1122 应收账款': 2036268933803311110,
    '1405 库存商品': 2036268933811699715,
    '2202 应付账款': 2036268933820088327,
    '3103 本年利润': 2036268933832671233,
    '5601 销售费用': 2036268933836865552,
    '5602 管理费用': 2036268933841059849,
}


def get_client():
    sess = load_session()
    client = WukongClient(sess['base_url'], token=sess['token'])
    client.post('/financeAccountSet/switchAccountSet', params={'accountId': ACCOUNT_ID})
    return client


def format_num(v):
    if v is None:
        return ''
    if isinstance(v, (int, float)):
        return f'{v:.2f}'
    return str(v)


def get_reports(client):
    balance_rows = report.balance_sheet(client, 1, PERIOD_END)
    income_rows = report.income_statement(client, 1, PERIOD_END)
    cash_rows = report.cash_flow_statement(client, 1, PERIOD_END)
    try:
        balance_check = report.balance_sheet_check(client, 1, PERIOD_END)
    except Exception as e:
        balance_check = {'error': str(e)}
    try:
        income_check = report.income_statement_check(client, 1, PERIOD_END)
    except Exception as e:
        income_check = {'error': str(e)}
    try:
        cash_check = report.cash_flow_check(client, 1, PERIOD_END)
    except Exception as e:
        cash_check = {'error': str(e)}
    return {
        'balance_sheet': {'rows': balance_rows, 'check': balance_check},
        'income_statement': {'rows': income_rows, 'check': income_check},
        'cash_flow_statement': {'rows': cash_rows, 'check': cash_check},
    }


def get_key_ledgers(client):
    result = {}
    for name, sid in KEY_SUBJECTS.items():
        rows = ledger.query_detail_account(client, sid, START_DATE, END_DATE)
        ending = rows[-1] if rows else {}
        business = [r for r in rows if r.get('type') == '2']
        result[name] = {
            'entry_count': len(business),
            'balance_direction': ending.get('balanceDirection'),
            'balance': ending.get('balance'),
            'debtor_balance': ending.get('debtorBalance'),
            'credit_balance': ending.get('creditBalance'),
        }
    return result


def pick_balance_rows(rows):
    return [r for r in rows if r.get('grade') in (1, 2) or abs(float(r.get('endPeriod') or 0)) > 1e-9]


def pick_income_rows(rows):
    return [r for r in rows if r.get('grade') in (1, 2) or abs(float(r.get('monthValue') or 0)) > 1e-9]


def pick_cash_rows(rows):
    return [r for r in rows if r.get('grade') in (1, 2) or abs(float(r.get('monthValue') or 0)) > 1e-9]


def write_balance_md(report_data):
    rows = pick_balance_rows(report_data['rows'])
    lines = [
        f'# {PERIOD_TEXT}资产负债表',
        '',
        f'- 单位：`{ACCOUNT_NAME}`',
        f'- 报表日期：`{PERIOD_END}`',
        '- 说明：以下为系统按当前凭证自动生成结果，尚需会计复核。',
        '',
        '| 项目 | 期初数 | 期末数 |',
        '| --- | ---: | ---: |',
    ]
    for row in rows:
        indent = '　' * max(int(row.get('grade', 1)) - 1, 0)
        lines.append(f"| {indent}{row.get('name','')} | {format_num(row.get('startPeriod'))} | {format_num(row.get('endPeriod'))} |")
    lines += ['', f"- 勾稽检查：`{report_data['check']}`"]
    BALANCE_MD.write_text('\n'.join(lines), encoding='utf-8')


def write_income_md(report_data):
    rows = pick_income_rows(report_data['rows'])
    lines = [
        f'# {PERIOD_TEXT}利润表',
        '',
        f'- 单位：`{ACCOUNT_NAME}`',
        f'- 报表日期：`{PERIOD_END}`',
        '- 说明：以下为系统按当前凭证自动生成结果，尚需会计复核。',
        '',
        '| 项目 | 本月数 | 本年累计 |',
        '| --- | ---: | ---: |',
    ]
    for row in rows:
        indent = '　' * max(int(row.get('grade', 1)) - 1, 0)
        lines.append(f"| {indent}{row.get('name','')} | {format_num(row.get('monthValue'))} | {format_num(row.get('yearValue'))} |")
    lines += ['', f"- 勾稽检查：`{report_data['check']}`"]
    INCOME_MD.write_text('\n'.join(lines), encoding='utf-8')


def write_cash_md(report_data):
    rows = pick_cash_rows(report_data['rows'])
    lines = [
        f'# {PERIOD_TEXT}现金流量表',
        '',
        f'- 单位：`{ACCOUNT_NAME}`',
        f'- 报表日期：`{PERIOD_END}`',
        '- 说明：以下为系统按当前凭证自动生成结果，尚需会计复核。',
        '',
        '| 项目 | 本月数 | 本年累计 |',
        '| --- | ---: | ---: |',
    ]
    for row in rows:
        indent = '　' * max(int(row.get('grade', 1)) - 1, 0)
        lines.append(f"| {indent}{row.get('name','')} | {format_num(row.get('monthValue'))} | {format_num(row.get('yearValue'))} |")
    lines += ['', f"- 勾稽检查：`{report_data['check']}`"]
    CASH_MD.write_text('\n'.join(lines), encoding='utf-8')


def write_ledger_md(ledger_data):
    lines = [
        f'# {PERIOD_TEXT}科目余额与关键科目结果',
        '',
        f'- 单位：`{ACCOUNT_NAME}`',
        f'- 期间：`{START_DATE}` 至 `{END_DATE}`',
        '- 说明：以下用于会计复核关键科目落账结果。',
        '',
        '| 科目 | 业务笔数 | 期末方向 | 期末余额 | 本期借方累计 | 本期贷方累计 |',
        '| --- | ---: | --- | ---: | ---: | ---: |',
    ]
    for name, item in ledger_data.items():
        lines.append(
            f"| {name} | {item['entry_count']} | {item['balance_direction']} | {format_num(item['balance'])} | {format_num(item['debtor_balance'])} | {format_num(item['credit_balance'])} |"
        )
    LEDGER_MD.write_text('\n'.join(lines), encoding='utf-8')


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def generate_docx(reports, ledgers):
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'PingFang SC'
    normal.font.size = Pt(10.5)

    add_title(doc, f'{PERIOD_TEXT}月度会计报表')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'单位：{ACCOUNT_NAME}')
    add_para(doc, '说明：以下为系统按当前凭证自动生成结果，尚需会计复核。')

    add_para(doc, '一、资产负债表', bold=True)
    add_table(doc, ['项目', '期初数', '期末数'], [
        ['　' * max(int(r.get('grade', 1)) - 1, 0) + str(r.get('name', '')), format_num(r.get('startPeriod')), format_num(r.get('endPeriod'))]
        for r in pick_balance_rows(reports['balance_sheet']['rows'])
    ])
    add_para(doc, f"勾稽检查：{reports['balance_sheet']['check']}")

    add_para(doc, '二、利润表', bold=True)
    add_table(doc, ['项目', '本月数', '本年累计'], [
        ['　' * max(int(r.get('grade', 1)) - 1, 0) + str(r.get('name', '')), format_num(r.get('monthValue')), format_num(r.get('yearValue'))]
        for r in pick_income_rows(reports['income_statement']['rows'])
    ])
    add_para(doc, f"勾稽检查：{reports['income_statement']['check']}")

    add_para(doc, '三、现金流量表', bold=True)
    add_table(doc, ['项目', '本月数', '本年累计'], [
        ['　' * max(int(r.get('grade', 1)) - 1, 0) + str(r.get('name', '')), format_num(r.get('monthValue')), format_num(r.get('yearValue'))]
        for r in pick_cash_rows(reports['cash_flow_statement']['rows'])
    ])
    add_para(doc, f"勾稽检查：{reports['cash_flow_statement']['check']}")

    add_para(doc, '四、关键科目结果', bold=True)
    add_table(doc, ['科目', '业务笔数', '期末方向', '期末余额', '本期借方累计', '本期贷方累计'], [
        [name, item['entry_count'], item['balance_direction'], format_num(item['balance']), format_num(item['debtor_balance']), format_num(item['credit_balance'])]
        for name, item in ledgers.items()
    ])

    doc.save(str(DOCX_PATH))


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    client = get_client()
    reports = get_reports(client)
    ledgers = get_key_ledgers(client)

    write_balance_md(reports['balance_sheet'])
    write_income_md(reports['income_statement'])
    write_cash_md(reports['cash_flow_statement'])
    write_ledger_md(ledgers)
    generate_docx(reports, ledgers)

    JSON_PATH.write_text(json.dumps({'reports': reports, 'ledgers': ledgers}, ensure_ascii=False, indent=2), encoding='utf-8')
    print(BALANCE_MD)
    print(INCOME_MD)
    print(CASH_MD)
    print(LEDGER_MD)
    print(DOCX_PATH)
    print(JSON_PATH)

if __name__ == '__main__':
    main()
